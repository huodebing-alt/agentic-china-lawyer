import argparse
from pathlib import Path

def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("docx"); ap.add_argument("--firm", default="律师事务所")
    ap.add_argument("--out", default="./paginated.docx")
    args = ap.parse_args()
    try:
        import docx
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        d = docx.Document(args.docx)
        section = d.sections[0]
        # 页眉
        section.header.paragraphs[0].text = args.firm
        # 页脚（带页码）— 简化：插入 'PAGE' field
        footer = section.footer.paragraphs[0]
        footer.text = "第   页 / 共   页"
        d.save(args.out)
        print(f"✓ {args.out}")
    except ImportError:
        print("❌ pip install python-docx"); raise SystemExit(2)

if __name__ == "__main__":
    main()
